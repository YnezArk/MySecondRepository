from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import re


def clean_html_tags(text):
    """清除HTML标签，保留纯文本内容"""
    if not text:
        return ""
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)


def split_json_objects(text):
    """将包含多个JSON对象的文本按大括号分割为单独的JSON字符串"""
    json_objects = []
    start = 0
    brace_count = 0

    for i, c in enumerate(text):
        if c == '{':
            brace_count += 1
            if brace_count == 1:
                start = i
        elif c == '}':
            brace_count -= 1
            if brace_count == 0:
                json_str = text[start:i + 1].strip()
                if json_str:
                    json_objects.append(json_str)

    return json_objects


def load_multi_json_file(file_path):
    """从单个文件加载多个JSON对象，返回数据列表"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 分割JSON对象并解析
        json_strings = split_json_objects(content)
        data_list = []
        for json_str in json_strings:
            try:
                data = json.loads(json_str)
                data_list.append(data)
            except json.JSONDecodeError as e:
                print(f"解析JSON对象失败: {e}")
                continue

        return data_list
    except Exception as e:
        print(f"加载文件 {file_path} 时出错：{str(e)}")
        return []


def match_answers_with_questions(questions_data, answers_data):
    """将题目数据与答案数据进行匹配"""
    matched_questions = {}

    # 建立题目ID到答案的映射
    answer_mapping = {}
    for qid, answer_info in answers_data['rt'].items():
        answer_ids = answer_info['answer'].split(',') if answer_info['answer'] else []
        answer_mapping[qid] = {
            'answer_ids': answer_ids,
            'score': answer_info['score']
        }

    # 遍历题目数据，匹配答案
    for part in questions_data['rt']['examBase']['workExamParts']:
        for question in part['questionDtos']:
            qid = str(question['id'])
            question_data = question.copy()

            if qid in answer_mapping:
                question_data['correct_answer'] = answer_mapping[qid]
                question_data['correct_letters'] = map_answer_ids_to_letters(
                    answer_mapping[qid]['answer_ids'],
                    question['questionOptions']
                )
                # 针对判断题，额外存储答案文本
                if question['questionType']['name'] == '判断题':
                    question_data['correct_text'] = get_judgment_answer_text(
                        answer_mapping[qid]['answer_ids'],
                        question['questionOptions']
                    )
            else:
                question_data['correct_answer'] = None
                question_data['correct_letters'] = []
                question_data['correct_text'] = ""

            matched_questions[qid] = question_data

    return matched_questions


def map_answer_ids_to_letters(answer_ids, options):
    """将选项ID映射为字母（A, B, C, D...）"""
    id_to_letter = {}
    for i, option in enumerate(options):
        id_to_letter[str(option['id'])] = chr(65 + i)

    correct_letters = []
    for aid in answer_ids:
        if aid in id_to_letter:
            correct_letters.append(id_to_letter[aid])

    return sorted(correct_letters)


def get_judgment_answer_text(answer_ids, options):
    """获取判断题的答案文本（对/错）"""
    for option in options:
        if str(option['id']) in answer_ids:
            return option['content']
    return ""


def export_chapter_to_doc(chapter_data, answers_data, doc, chapter_index):
    """将单个章节的题目添加到文档中"""
    exam_info = chapter_data['rt']['examBase']

    # 添加章节标题
    chapter_title = doc.add_heading(level=0)
    chapter_title_run = chapter_title.add_run(f"{exam_info['toChapter']} - {exam_info['name']}")
    chapter_title_run.font.size = Pt(14)
    chapter_title_run.bold = True
    chapter_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加章节基本信息
    info_para = doc.add_paragraph()
    info_run = info_para.add_run(f"课程名称：{exam_info['courseName']} | ")
    info_run = info_para.add_run(f"题目数量：{exam_info['problemNum']}题 | ")
    info_run = info_para.add_run(f"总分：{exam_info['totalScore']}分")
    doc.add_paragraph()  # 空行

    # 匹配题目和答案
    matched_questions = match_answers_with_questions(chapter_data, answers_data)

    # 遍历所有题目
    total_questions = 0

    for part_index, part in enumerate(exam_info['workExamParts'], 1):
        if len(exam_info['workExamParts']) > 1:
            part_title = doc.add_heading(level=1)
            part_title.add_run(f"第{part_index}部分（共{part['questionCount']}题）")
            doc.add_paragraph()

        for question in part['questionDtos']:
            total_questions += 1
            qid = str(question['id'])
            matched_question = matched_questions[qid]

            # 添加题目编号和类型
            type_para = doc.add_paragraph()
            type_text = f"第{total_questions}题：{question['questionType']['name']}（{question['questionScore']}分）"

            if matched_question['correct_answer']:
                score = matched_question['correct_answer']['score']
                type_text += f" - 得分：{score}分"

            type_run = type_para.add_run(type_text)
            type_run.bold = True
            type_run.font.size = Pt(11)

            # 添加题目内容
            question_text = clean_html_tags(question['name']).strip()
            content_para = doc.add_paragraph()
            content_para.paragraph_format.first_line_indent = Pt(24)
            content_para.add_run(question_text)

            # 添加选项（选择题）
            if question['questionOptions'] and question['questionType']['name'] in ['多选题', '单选题']:
                for option in question['questionOptions']:
                    option_para = doc.add_paragraph()
                    option_para.paragraph_format.left_indent = Pt(36)
                    option_para.paragraph_format.space_after = Pt(6)

                    option_prefix = f"{chr(65 + question['questionOptions'].index(option))}."
                    option_text = f"{option_prefix} {option['content']}"

                    if str(option['id']) in matched_question['correct_answer']['answer_ids']:
                        option_run = option_para.add_run(option_text)
                        option_run.bold = True
                        option_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                    else:
                        option_para.add_run(option_text)

            # 判断题处理
            elif question['questionType']['name'] == '判断题':
                judgment_para = doc.add_paragraph()
                judgment_para.paragraph_format.left_indent = Pt(36)

                for option in question['questionOptions']:
                    option_text = f"{option['content']}  "
                    if str(option['id']) in matched_question['correct_answer']['answer_ids']:
                        option_run = judgment_para.add_run(option_text)
                        option_run.bold = True
                        option_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                    else:
                        judgment_para.add_run(option_text)

            # 添加正确答案提示
            answer_para = doc.add_paragraph()
            answer_para.paragraph_format.left_indent = Pt(36)

            if question['questionType']['name'] == '判断题':
                # 判断题显示"对/错"
                answer_text = f"正确答案：{matched_question['correct_text']}"
            else:
                # 其他题型显示选项字母
                answer_text = "正确答案："
                if matched_question['correct_letters']:
                    answer_text += " " + ", ".join(matched_question['correct_letters'])
                else:
                    answer_text += " 无"

            answer_run = answer_para.add_run(answer_text)
            answer_run.bold = True
            answer_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

            doc.add_paragraph()  # 题目间空行

    return total_questions


def export_multi_chapter_exam(questions_file, answers_file, filename="完整试题集（含答案）.docx"):
    """从单个题目文件和单个答案文件导出多章节试题"""
    # 加载所有章节的题目和答案数据
    chapters_data = load_multi_json_file(questions_file)
    answers_data_list = load_multi_json_file(answers_file)

    if len(chapters_data) != len(answers_data_list):
        print(f"警告：题目章节数（{len(chapters_data)}）与答案章节数（{len(answers_data_list)}）不匹配！")
        return None

    # 创建Word文档
    doc = Document()

    # 设置文档全局字体
    style = doc.styles['Normal']
    style.font.name = u'宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    style.font.size = Pt(10.5)

    # 添加总标题
    total_title = doc.add_heading(level=0)
    if chapters_data:
        first_chapter = chapters_data[0]
        course_name = first_chapter['rt']['examBase']['courseName'] if first_chapter else "试题集"
        total_title_run = total_title.add_run(f"{course_name} 完整试题集（含答案）")
    else:
        total_title_run = total_title.add_run("完整试题集（含答案）")
    total_title_run.font.size = Pt(18)
    total_title_run.bold = True
    total_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    doc.add_paragraph()

    total_questions = 0

    # 逐个导入章节
    for i, (chapter_data, answer_data) in enumerate(zip(chapters_data, answers_data_list), 1):
        print(f"正在导入第{i}章...")

        # 将章节添加到文档
        chapter_questions = export_chapter_to_doc(chapter_data, answer_data, doc, i)
        total_questions += chapter_questions

        # 章节之间添加分页
        if i < len(chapters_data):
            doc.add_page_break()

    # 添加统计信息
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    stats_text = f"试题集统计：共{len(chapters_data)}章，总题目数{total_questions}题"
    stats_run = stats_para.add_run(stats_text)
    stats_run.italic = True
    stats_run.font.size = Pt(12)
    stats_run.bold = True

    # 保存文档
    doc.save(filename)
    print(f"\n✅ 多章节试题集导出成功！")
    print(f"📁 文件名称：{filename}")
    print(f"📊 总章节数：{len(chapters_data)}章")
    print(f"📊 总题目数：{total_questions}题")

    return filename


# 执行导出
if __name__ == "__main__":
    # 只需指定两个文件：一个包含所有章节的题目，一个包含所有章节的答案
    questions_file = "questions_data.txt"  # 所有章节题目数据（多个JSON对象）
    answers_file = "answers_data.txt"  # 所有章节答案数据（多个JSON对象）

    # 导出为单个文档
    export_multi_chapter_exam(questions_file, answers_file)